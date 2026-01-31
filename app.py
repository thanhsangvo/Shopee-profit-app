import streamlit as st
import plotly.graph_objects as go
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit.components.v1 as components

# 1. Khởi tạo session_state để lưu danh sách sản phẩm
if 'danh_sach_sp' not in st.session_state:
    st.session_state.danh_sach_sp = []
    
# 1. Cấu hình trang & Giao diện
st.set_page_config(page_title="Shopee Profit 2026", page_icon="♥️", layout="wide")

# Thay G-XXXXXXXXXX bằng ID thực tế của bạn
GA_ID = "G-X11FLFF1S7"

ga_code = f"""
    <script async src="https://www.googlesyndication.com/pagead/js/adsbygoogle.js?client={GA_ID}" crossorigin="anonymous"></script>
    <script async src="https://www.googletagmanager.com/gtag/js?id={GA_ID}"></script>
    <script>
        window.dataLayer = window.dataLayer || [];
        function gtag(){{dataLayer.push(arguments);}}
        gtag('js', new Date());
        gtag('config', '{GA_ID}');
    </script>
"""
# Chèn vào app (thường đặt ở đầu hoặc cuối file)
components.html(ga_code, height=0)

# Hàm định dạng tiền tệ chuẩn Việt Nam
def format_vnd(amount):
    return f"{amount:,.0f}".replace(",", ".") + " đ"

# CSS tùy chỉnh giao diện
st.markdown("""
    <style>
    .stMetric { background-color: #ffffff; padding: 20px; border-radius: 12px; border: 1px solid #ff4d2d; }
    .main { background-color: #fafafa; }
    /* Tùy chỉnh ô Metric cho cả sáng và tối */
    [data-testid="stMetric"] {
        background-color: rgba(255, 77, 45, 0.05);
        border: 1px solid rgba(255, 77, 45, 0.3);
        padding: 15px;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    [data-testid="stMetricLabel"] {
        color: #ff4d2d !important;
        font-weight: bold !important;
    }
    [data-testid="stMetricValue"] {
    color: #2ecc71 !important; 
    }
    .hoavon-box {
        background-color: rgba(0, 104, 201, 0.05);
        border: 1px dashed #0068c9;
        padding: 10px;
        border-radius: 10px;
        margin-top: 10px;
        margin-bottom: 10px;
    }
    
    div.stButton > button:first-child {
        background-color: #0068c9;
        color: white;
        border: none;
        transition: all 0.3s ease;
    }

    /* Hiệu ứng khi di chuột qua (hover) */
    div.stButton > button:first-child:hover {
        background-color: #0056a3;
        border: none;
        color: white;
        transform: scale(1.02);
    }
    
   /* Nút Xóa danh sách (Nút thứ 2 trong trang hoặc dựa trên text) */
    button[kind="secondary"]:has(div:contains("Xóa danh sách")) {
        background-color: #ff4b4b;
        color: white;
    }
    
    /* Style cho Banner có chứa ảnh sản phẩm */
    .affiliate-main-container {
        background: linear-gradient(135deg, #ff4d2d 0%, #ff8e3c 100%);
        padding: 15px;
        border-radius: 15px;
        color: white;
        display: flex;
        align-items: center;
        gap: 20px;
        box-shadow: 0 4px 15px rgba(255, 77, 45, 0.2);
        margin: 15px 0;
    }

    .affiliate-image {
        width: 80px;
        height: 80px;
        object-fit: cover;
        border-radius: 10px;
        border: 2px solid white;
        flex-shrink: 0;
    }

    .affiliate-content {
        flex-grow: 1;
    }

    @media (max-width: 768px) {
        .affiliate-main-container {
            flex-direction: column;
            text-align: center;
        }
        .affiliate-image {
            width: 120px;
            height: 120px;
        }
        .affiliate-main-button {
            width: 100%;
        }
    }
    </style>
    """, unsafe_allow_html=True)

st.title("Shopee Profit Master (Cập nhật 2026)")
st.caption("Dữ liệu phí cố định trích xuất từ chính sách mới nhất (29/12/2025)")

# 2. Thanh bên (Sidebar) - Nhập liệu thông số cố định
st.sidebar.header("⚙️ Cài đặt chi phí sàn")
phi_thanh_toan = st.sidebar.number_input("Phí thanh toán (%)", value=4.91, step=0.01, help="Mặc định 4.91% theo quy định mới")
thue_tncn = st.sidebar.number_input("Thuế TNCN & GTGT (%)", value=1.5, step=0.1, help="Thường là 1.5% doanh thu cho hộ kinh doanh")

# Link ảnh sản phẩm thực tế từ Shopee
url_anh_sp = "https://down-vn.img.susercontent.com/file/vn-11134207-820l4-mhurso4d84xv4e.webp" # Đây là ví dụ ảnh máy in
url_sp = "https://s.shopee.vn/1gCmjbDbTa"
affiliate_main_html = f"""
<div class="affiliate-main-container">
    <img src="{url_anh_sp}" class="affiliate-image">
    <div class="affiliate-content">
        <div style="font-size: 1.1rem; font-weight: bold;">🎁 ƯU ĐÃI: Máy In Đơn AYIN Bluetooth</div>
        <div style="font-size: 0.85rem; opacity: 0.9; margin-top: 5px;">
            Giảm ngay 20% + Tặng giấy in cho chủ shop dùng App!
        </div>
    </div>
    <a href="{url_sp}" target="_blank" class="affiliate-main-button">
        🛒 Mua Ngay
    </a>
</div>
"""
st.markdown(affiliate_main_html, unsafe_allow_html=True)

# 3. Thân trang - Nhập liệu sản phẩm
col_input, col_result = st.columns([1, 1])

with col_input:
    st.subheader("📦 Thông tin sản phẩm")
    ten_sp = st.text_input("Tên sản phẩm", "Sản phẩm A")
    
    c1, c2 = st.columns(2)
    with c1:
        gia_von = st.number_input("Giá vốn (VNĐ)", min_value=0, value=120000, step=1000, format="%d")
        st.caption(f"Xác nhận: :blue[{format_vnd(gia_von)}]")
        
        gia_ban = st.number_input("Giá bán (VNĐ)", min_value=0, value=200000, step=1000, format="%d")
        st.caption(f"Xác nhận: :blue[{format_vnd(gia_ban)}]")
        
    with c2:
        dict_nganh_hang = {
            "Thời trang Nữ/Nam/Trẻ em (13.5%)": 13.5,
            "Sắc đẹp - Chăm sóc da mặt (14.0%)": 14.0,
            "Sức khỏe - Thực phẩm chức năng (14.0%)": 14.0,
            "Điện thoại & Phụ kiện (12.0%)": 12.0,
            "Thiết bị âm thanh/Cameras (10.0%)": 10.0,
            "Máy tính & Laptop (Linh kiện) (7.5%)": 7.5,
            "Điện tử - Tivi & Phụ kiện (8.0%)": 8.0,
            "Điện gia dụng lớn (7.5%)": 7.5,
            "Đồ gia dụng nhà bếp (10.0%)": 10.0,
            "Thực phẩm & Đồ uống (11.0%)": 11.0,
            "Chăm sóc thú cưng (13.0%)": 13.0,
            "Ô tô - Phụ tùng & Chăm sóc (13.0%)": 13.0,
            "Voucher & Dịch vụ (11.0%)": 11.0,
            "Laptop / Màn hình / Điện thoại (máy) (2.0%)": 2.0
        }
        ten_nganh = st.selectbox("Ngành hàng (Phí cố định mới)", options=list(dict_nganh_hang.keys()))
        phi_nganh_hang = dict_nganh_hang[ten_nganh]

    st.subheader("🚀 Gói dịch vụ tham gia")
    g1, g2 = st.columns(2)
    with g1:
        fsx = st.checkbox("Freeship Xtra (7%)", value=True)
        hxx = st.checkbox("Hoàn Xu Xtra (5%)", value=False)
    with g2:
        phi_bao_bi = st.number_input("Phí bao bì/ đóng gói", value=2000, step=500, format="%d")
        st.caption(f"Xác nhận: :blue[{format_vnd(phi_bao_bi)}]")
        
        phi_ads = st.number_input("Phí Marketing/Ads dự tính", value=5000, step=500, format="%d")
        st.caption(f"Xác nhận: :blue[{format_vnd(phi_ads)}]")

# 4. Logic Tính Toán
phi_ship_shopee = 0 
tien_phi_thanh_toan = gia_ban * (phi_thanh_toan / 100)
tien_phi_co_dinh = gia_ban * (phi_nganh_hang / 100)
tien_fsx = min(gia_ban * 0.07, 40000) if fsx else 0
tien_hxx = min(gia_ban * 0.05, 20000) if hxx else 0
tien_thue = gia_ban * (thue_tncn / 100)

tong_phi_san = tien_phi_thanh_toan + tien_phi_co_dinh + tien_fsx + tien_hxx
tong_chi_phi = gia_von + tong_phi_san + tien_thue + phi_bao_bi + phi_ads
loi_nhuan = gia_ban - tong_chi_phi
bien_ln = (loi_nhuan / gia_ban * 100) if gia_ban > 0 else 0

# Tính giá hòa vốn
tong_phi_pct = phi_thanh_toan + phi_nganh_hang + thue_tncn + (7.0 if fsx else 0) + (5.0 if hxx else 0)
gia_hoa_von = (gia_von + phi_bao_bi + phi_ads) / (1 - tong_phi_pct/100) if tong_phi_pct < 100 else 0

# 5. Hiển thị kết quả
with col_result:
    st.subheader("📊 Phân tích lợi nhuận")

    res1, res2 = st.columns(2)
    with res1:
        st.metric(label="Lợi nhuận ròng", value=format_vnd(loi_nhuan), delta=f"{bien_ln:.1f}% (Biên LN)", delta_color="normal" if loi_nhuan > 0 else "inverse")
    with res2:
        st.metric(label="Tổng phí sàn", value=format_vnd(tong_phi_san), delta="Đã gồm thuế", delta_color="off")
        
    st.markdown('<div class="hoavon-box">', unsafe_allow_html=True)
    st.metric(label="🎯 Giá bán tối thiểu để hòa vốn", value=format_vnd(gia_hoa_von), help="Bán dưới giá này bạn sẽ bị lỗ.")
    
    if gia_ban < gia_hoa_von and gia_ban > 0:
        st.error(f"⚠️ Cần tăng giá thêm ít nhất {format_vnd(gia_hoa_von - gia_ban)} để không bị lỗ.")
    elif gia_ban > 0:
        st.success(f"💎 Bạn đang bán cao hơn giá hòa vốn {format_vnd(gia_ban - gia_hoa_von)}.")
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Vẽ biểu đồ
    labels = ['Giá vốn', 'Phí sàn', 'Thuế', 'Vận hành', 'Lợi nhuận']
    values = [gia_von, tong_phi_san, tien_thue, (phi_bao_bi + phi_ads), max(0, loi_nhuan)]
    colors = ['#3498db', '#ff4d2d', '#95a5a6', '#f1c40f', '#2ecc71']

    fig = go.Figure(data=[go.Pie(labels=labels, values=values, hole=.4, marker=dict(colors=colors, line=dict(color='#ffffff', width=1)), textinfo='percent')])
    fig.update_layout(margin=dict(t=30, b=0, l=0, r=0), height=350, paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', legend=dict(orientation="h", yanchor="bottom", y=-0.2, xanchor="center", x=0.5, font=dict(color="gray")), annotations=[dict(text='Cơ cấu', x=0.5, y=0.5, font_size=18, showarrow=False, font=dict(color="gray"))])
    st.plotly_chart(fig, use_container_width=True)

    if loi_nhuan < 0:
        st.error(f"⚠️ Đơn hàng này đang lỗ: {format_vnd(loi_nhuan)}")
    elif bien_ln < 15:
        st.warning("⚠️ Biên lợi nhuận mỏng (dưới 15%), hãy cẩn thận.")
    else:
        st.success("✅ Chỉ số lợi nhuận rất tốt!")

# 6. Tính năng Cảnh báo an toàn (Risk Management)
    st.markdown("---")
    st.subheader("🛡️ Đánh giá rủi ro")
    
    # Tính toán các chỉ số rủi ro
    ty_trong_phi_san = (tong_phi_san / gia_ban * 100) if gia_ban > 0 else 0
    diem_rui_ro = 0
    loi_khuyen = []

    # Kiểm tra tỷ trọng phí sàn
    if ty_trong_phi_san > 25:
        diem_rui_ro += 2
        loi_khuyen.append("- ⚠️ **Phí sàn quá cao:** Chi phí sàn chiếm >25% giá bán. Bạn nên xem xét lại việc tham gia cùng lúc quá nhiều gói Xtra hoặc tăng giá bán.")
    elif ty_trong_phi_san > 18:
        loi_khuyen.append("- ℹ️ **Phí sàn trung bình:** Mức phí này khá phổ biến khi tham gia đầy đủ các gói dịch vụ.")

    # Kiểm tra biên lợi nhuận
    if bien_ln < 10 and bien_ln > 0:
        diem_rui_ro += 1
        loi_khuyen.append("- ⚠️ **Biên lãi mỏng:** Chỉ cần khách trả hàng hoặc chạy Ads quá tay là bạn sẽ lỗ.")
    elif bien_ln <= 0:
        diem_rui_ro += 3
        loi_khuyen.append("- 🚨 **BÁO ĐỘNG ĐỎ:** Bạn đang bán lỗ! Hãy điều chỉnh giá bán hoặc giá vốn ngay lập tức.")

    # Hiển thị mức độ an toàn
    if diem_rui_ro == 0:
        st.success("✅ **Mức độ an toàn: CAO.** Chỉ số tài chính của sản phẩm này rất bền vững.")
    elif diem_rui_ro <= 2:
        st.warning("⚠️ **Mức độ an toàn: TRUNG BÌNH.** Cần tối ưu thêm chi phí vận hành.")
    else:
        st.error("🔥 **Mức độ an toàn: THẤP.** Cần thay đổi chiến lược ngay.")

    for lk in loi_khuyen:
        st.write(lk)

# 5. Nút LƯU VÀO DANH SÁCH (Đặt cuối phần col_input)
with col_input:
    if st.button("➕ Lưu vào danh sách so sánh"):
        new_data = {
            "Tên SP": ten_sp,
            "Giá vốn": gia_von,
            "Giá bán": gia_ban,
            "Phí sàn": tong_phi_san,
            "Lợi nhuận": loi_nhuan,
            "Biên LN (%)": round(bien_ln, 2),
            "Giá hòa vốn": round(gia_hoa_von, 0)
        }
        st.session_state.danh_sach_sp.append(new_data)
        st.toast(f"Đã lưu {ten_sp} vào danh sách!", icon="✅")
        
# 6. HIỂN THỊ BẢNG SO SÁNH Ở CUỐI TRANG
st.markdown("---")
st.subheader("📋 Danh sách so sánh & Tổng hợp")

if st.session_state.danh_sach_sp:
    df = pd.DataFrame(st.session_state.danh_sach_sp)
    
    # --- TÍNH TOÁN TỔNG CỘNG ---
    tong_ln_danh_sach = df["Lợi nhuận"].sum()
    so_luong_sp = len(df)
    bien_ln_tb = df["Biên LN (%)"].mean()

    # Hiển thị các con số tổng quát lên trên bảng
    t1, t2, t3 = st.columns(3)
    t1.metric("Tổng số mặt hàng", f"{so_luong_sp} SP")
    t2.metric("Tổng lợi nhuận dự tính", format_vnd(tong_ln_danh_sach))
    t3.metric("Biên LN trung bình", f"{bien_ln_tb:.1f}%")

    # --- ĐỊNH DẠNG VÀ CĂN LỀ BẢNG ---
    # Sử dụng column_config để căn lề giữa (center) và định dạng số
    st.dataframe(
        df, 
        use_container_width=True,
        hide_index=True,
        column_config={
            "Tên SP": st.column_config.TextColumn("Tên sản phẩm", width="medium"),
            "Giá vốn": st.column_config.NumberColumn("Giá vốn", format="%d", help="Đơn vị: VNĐ"),
            "Giá bán": st.column_config.NumberColumn("Giá bán", format="%d"),
            "Phí sàn": st.column_config.NumberColumn("Phí sàn", format="%d"),
            "Lợi nhuận": st.column_config.NumberColumn("Lợi nhuận", format="%d"),
            "Biên LN (%)": st.column_config.NumberColumn("Biên LN (%)", format="%.2f%%"),
            "Giá hòa vốn": st.column_config.NumberColumn("Giá hòa vốn", format="%d"),
        }
    )

    # CSS để ép căn giữa nội dung trong bảng (Streamlit hiện tại căn lề dựa trên kiểu dữ liệu, 
    # nhưng chúng ta có thể bổ sung CSS để giao diện đồng nhất hơn)
    st.markdown("""
        <style>
            [data-testid="stTable"] td { text-align: center !important; }
            [data-testid="stDataFrame"] div[data-testid="stVerticalBlock"] > div { text-align: center; }
        </style>
    """, unsafe_allow_html=True)

    # --- CÁC NÚT THAO TÁC (Xóa & Tải Excel) ---
    c1, c2 = st.columns([1, 5])
    with c1:
        if st.button("🗑️ Xóa danh sách"):
            st.session_state.danh_sach_sp = []
            st.rerun()

    with c2:
        output = BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='Lợi_Nhuận_Shopee')
        
        st.download_button(
            label="📥 Tải về file Excel (.xlsx)",
            data=output.getvalue(),
            file_name="tong_hop_loi_nhuan_shopee.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
else:
    st.info("Chưa có sản phẩm nào được lưu. Hãy nhập thông tin và nhấn 'Lưu vào danh sách'.")
    
st.markdown("---")
st.caption("Dữ liệu trích xuất từ phụ lục Phí cố định áp dụng từ 29/12/2025 (đã bao gồm thuế GTGT 8%).")

# --- Hàm tạo file DOCX ---
def create_docx(ten_sp, gia_ban, gia_von, loi_nhuan, bien_ln, gia_hoa_von, tong_phi_san, tien_thue, phi_bao_bi, phi_ads):
    doc = Document()
    
    # Tiêu đề báo cáo
    title = doc.add_heading('BÁO CÁO PHÂN TÍCH LỢI NHUẬN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Tên sản phẩm: {ten_sp}")
    doc.add_paragraph(f"Ngày lập báo cáo: {pd.Timestamp.now().strftime('%d/%m/%Y')}")

    # Phần 1: Tóm tắt tài chính
    doc.add_heading('1. Chỉ số tài chính chính', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Chỉ số'
    hdr_cells[1].text = 'Giá trị'

    data = [
        ("Giá bán", format_vnd(gia_ban)),
        ("Giá vốn", format_vnd(gia_von)),
        ("Lợi nhuận ròng", format_vnd(loi_nhuan)),
        ("Biên lợi nhuận", f"{bien_ln:.2f}%"),
        ("Giá hòa vốn tối thiểu", format_vnd(gia_hoa_von))
    ]

    for item, value in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = value

    # Phần 2: Chi tiết chi phí
    doc.add_heading('2. Chi tiết các loại chi phí', level=1)
    phi_para = doc.add_paragraph()
    phi_para.add_run(f"- Tổng phí sàn (Chưa thuế 8%): {format_vnd(tong_phi_san - tien_thue)}\n")
    phi_para.add_run(f"- Tổng chi phí sàn: {format_vnd(tong_phi_san)}\n")
    phi_para.add_run(f"- Thuế TNCN & GTGT (1.5%): {format_vnd(tien_thue)}\n")
    phi_para.add_run(f"- Phí đóng gói: {format_vnd(phi_bao_bi)}\n")
    phi_para.add_run(f"- Phí Marketing/Ads: {format_vnd(phi_ads)}")

    # Lời kết
    doc.add_paragraph('\n---')
    footer = doc.add_paragraph('Báo cáo được tạo tự động bởi Shopee Profit Master.')
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Lưu vào bộ nhớ đệm để tải về
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- Giao diện nút bấm trên App ---
st.markdown("---")
st.subheader("📄 Xuất báo cáo chuyên nghiệp")

docx_file = create_docx(
    ten_sp, gia_ban, gia_von, loi_nhuan, bien_ln, gia_hoa_von, 
    tong_phi_san, tien_thue, phi_bao_bi, phi_ads
)

st.download_button(
    label="📥 Tải báo cáo Word (.docx)",
    data=docx_file,
    file_name=f"Bao_cao_{ten_sp.replace(' ', '_')}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
st.info("Để xuất báo cáo dưới dạng PDF hoặc Ảnh: Nhấn `Ctrl + P` (Windows) hoặc `Cmd + P` (Mac) trên trình duyệt, sau đó chọn 'Lưu dưới dạng PDF' hoặc 'In màn hình'.")
st.markdown("<p style='color: #ff4d2d; font-weight: bold; text-align: center;'>⚠️ Lưu ý: Kết quả mang tính chất tham khảo. Luôn kiểm tra đối soát thực tế trên Kênh Người Bán.</p>", unsafe_allow_html=True)